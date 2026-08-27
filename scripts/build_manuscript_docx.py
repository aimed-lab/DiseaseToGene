#!/usr/bin/env python3
"""Build the submission .docx from the v2 markdown sources.

    python scripts/build_manuscript_docx.py

Writes manuscript/v2/Modality_Fit_Manuscript_v2.docx and
       manuscript/v2/Modality_Fit_Supplementary_v2.docx

The markdown files are the single source of truth. The v1 .docx had diverged from the v1
.md — it carried a related-work paragraph and 15 references the .md never had — which is
the failure this script exists to prevent: regenerate rather than hand-edit.

Run it AFTER scripts/verifyManuscriptClaims.ts passes, so the document can only be built
from numbers that have been checked.
"""
import os, re, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
V2 = os.path.join(ROOT, "manuscript", "v2")
FIGDIR = os.path.join(V2, "figures")

BODY_FONT, HEAD_FONT, MONO_FONT = "Times New Roman", "Calibri", "Consolas"
INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTE = RGBColor(0x59, 0x59, 0x59)
ACCENT = RGBColor(0x1C, 0x44, 0x66)

FIGURES = [
    ("Figure 1", "figure1_goal_conditioning.png", 6.6),
    ("Figure 2", "figure2_exclusion.png", 5.4),
    ("Figure 3", "figure3_permissiveness.png", 6.0),
    ("Figure 4", "figure4_ablation.png", 6.0),
    ("Figure 5", "figure5_goal_blind_baseline.png", 6.0),
]


def shade(cell, hexcolor):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear"); el.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)


def setup(doc):
    s = doc.sections[0]
    s.page_width, s.page_height = Inches(8.5), Inches(11)
    for m in ("left_margin", "right_margin"):
        setattr(s, m, Inches(1.0))
    s.top_margin = s.bottom_margin = Inches(1.0)
    n = doc.styles["Normal"]
    n.font.name = BODY_FONT; n.font.size = Pt(11); n.font.color.rgb = INK
    n.paragraph_format.space_after = Pt(8)
    n.paragraph_format.line_spacing = 1.15
    for lvl, size in ((1, 16), (2, 13), (3, 11.5)):
        h = doc.styles[f"Heading {lvl}"]
        h.font.name = HEAD_FONT; h.font.size = Pt(size); h.font.bold = True
        h.font.color.rgb = ACCENT if lvl < 3 else INK
        h.paragraph_format.space_before = Pt(16 if lvl == 1 else 12)
        h.paragraph_format.space_after = Pt(5)
        h.paragraph_format.keep_with_next = True


# ── inline markdown: **bold**, *italic*, `code`, ^superscript ────────────────
# Recursive, because these nest: the abstract's word-count line is an italic span that
# contains a code span, and a flat pass renders the backticks literally.
TOKEN = re.compile(r"(\*\*.+?\*\*|(?<!\*)\*(?!\*)[^*]+\*(?!\*)|`[^`]+`|\^[0-9]+(?:,[0-9]+)*)")
ESC = "\x01"  # placeholder for a backslash-escaped asterisk, restored at the leaves


def emit(par, text, bold=False, italic=False):
    text = text.replace(r"\*", ESC).replace(r"\_", "_").replace(r"\#", "#")
    for piece in TOKEN.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**") and len(piece) > 4:
            emit(par, piece[2:-2], bold=True, italic=italic)
        elif piece.startswith("`") and piece.endswith("`"):
            r = par.add_run(piece[1:-1].replace(ESC, "*"))
            r.font.name = MONO_FONT; r.font.size = Pt(9.5)
            r.bold, r.italic = bold, italic
        elif piece.startswith("^"):
            r = par.add_run(piece[1:]); r.font.superscript = True
            r.bold, r.italic = bold, italic
        elif len(piece) > 2 and piece[0] == "*" and piece[-1] == "*":
            emit(par, piece[1:-1], bold=bold, italic=True)
        else:
            r = par.add_run(piece.replace(ESC, "*"))
            r.bold, r.italic = bold, italic


def add_table(doc, rows):
    header, body = rows[0], rows[1:]
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(header):
        c = t.rows[0].cells[i]; c.text = ""
        p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(2)
        emit(p, h)
        for r in p.runs:
            r.bold = True; r.font.size = Pt(9); r.font.name = HEAD_FONT
        shade(c, "E8EDF2")
    for row in body:
        cells = t.add_row().cells
        for i, val in enumerate(row[:len(header)]):
            cells[i].text = ""
            p = cells[i].paragraphs[0]; p.paragraph_format.space_after = Pt(2)
            emit(p, val)
            for r in p.runs:
                r.font.size = Pt(9); r.font.name = BODY_FONT
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def render(doc, md, *, with_figures=False):
    lines = md.split("\n")
    i, in_title = 0, True
    while i < len(lines):
        ln = lines[i].rstrip()

        # table
        if ln.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s:\-|]+\|$", lines[i + 1].strip()):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not re.match(r"^[\s:\-]+$", "".join(cells)):
                    rows.append(cells)
                i += 1
            add_table(doc, rows)
            continue

        if not ln.strip():
            i += 1; continue

        if ln.startswith("---") and set(ln.strip()) == {"-"}:
            i += 1; continue

        # headings
        m = re.match(r"^(#{1,4})\s+(.*)", ln)
        if m:
            level, txt = len(m.group(1)), m.group(2)
            if level == 1 and in_title:
                p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(14)
                r = p.add_run(txt); r.bold = True; r.font.size = Pt(17)
                r.font.name = HEAD_FONT; r.font.color.rgb = ACCENT
                in_title = False
            else:
                # Markdown '#' is the document title, so '##' sections must become Word
                # Heading 1 — otherwise the navigation pane and any generated TOC are
                # indented one level too deep and show no top-level entries.
                doc.add_heading(re.sub(r"\*\*", "", txt), level=max(1, level - 1))
            i += 1; continue

        # blockquote (may span lines)
        if ln.startswith(">"):
            buf = []
            while i < len(lines) and lines[i].startswith(">"):
                buf.append(lines[i].lstrip(">").strip()); i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(10)
            emit(p, " ".join(x for x in buf if x))
            for r in p.runs:
                r.font.size = Pt(10); r.font.color.rgb = MUTE
            continue

        # list item
        m = re.match(r"^(\s*)([-*]|\d+\.)\s+(.*)", ln)
        if m:
            indent, marker, txt = m.group(1), m.group(2), m.group(3)
            txt = re.sub(r"^\[[ xX]\]\s*", "", txt)
            style = "List Number" if marker[0].isdigit() else "List Bullet"
            p = doc.add_paragraph(style=style)
            p.paragraph_format.left_indent = Inches(0.3 + 0.25 * (len(indent) // 2))
            p.paragraph_format.space_after = Pt(3)
            emit(p, txt)
            i += 1; continue

        # paragraph
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        emit(p, ln)
        i += 1

    if with_figures:
        doc.add_page_break()
        doc.add_heading("Figures", level=1)
        caps = dict(re.findall(r"^- \*\*(Figure \d)\.\*\*\s*(.+)$", md, re.M))
        for name, fn, width in FIGURES:
            path = os.path.join(FIGDIR, fn)
            if not os.path.exists(path):
                print(f"  ! missing {fn}"); continue
            doc.add_picture(path, width=Inches(width))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp = doc.add_paragraph()
            cp.paragraph_format.space_after = Pt(18)
            r = cp.add_run(f"{name}. "); r.bold = True; r.font.size = Pt(9.5)
            body = re.sub(r"\*(.+?)\*", r"\1", caps.get(name, ""))
            body = re.sub(r"`([^`]+)`", r"\1", body)
            r2 = cp.add_run(body); r2.font.size = Pt(9.5); r2.font.color.rgb = MUTE
            print(f"  + {name}")


def build(src, out, title, with_figures):
    with open(os.path.join(V2, src), encoding="utf-8") as f:
        md = f.read()
    doc = Document()
    setup(doc)
    doc.core_properties.title = title
    doc.core_properties.comments = f"Generated from manuscript/v2/{src} by scripts/build_manuscript_docx.py"
    print(f"{src} -> {os.path.basename(out)}")
    render(doc, md, with_figures=with_figures)
    doc.save(out)
    print(f"  wrote {out}\n")


if __name__ == "__main__":
    build("manuscript_v2.md", os.path.join(V2, "Modality_Fit_Manuscript_v2.docx"),
          "Mechanistic goal changes which modality can engage a target", True)
    build("supplementary_v2.md", os.path.join(V2, "Modality_Fit_Supplementary_v2.docx"),
          "Supplementary Information (v2)", False)
